"""
Parte 1 — Módulo: src/infrastructure/ingestion/file_extractor.py

ESCOPO
------
Função extract(content: bytes, filename: str) -> list[dict[str, Any]]
Detecta o tipo pelo nome (extensão) e extrai texto/registros.

Formatos suportados: .txt, .csv, .json, .docx, .xlsx, .pdf
Formato não suportado: levanta ValueError.

TIPOS DE TESTE
--------------
- unit         : todas as funções (sem rede, sem disco persistente)
- regression   : extensão desconhecida (.pptx) deve levantar
"""
import io
import json

import pytest

from tests.shared.log_helper import logged

from src.infrastructure.ingestion.file_extractor import extract


@pytest.mark.unit
@pytest.mark.part1
class TestExtractTxt:
    """Arquivos .txt — split por parágrafo (duplo newline)."""

    @logged
    def test_splits_by_double_newline(self):
        content = b"Primeiro paragrafo.\n\nSegundo paragrafo."
        records = extract(content, "doc.txt")
        assert len(records) == 2
        assert records[0]["text"] == "Primeiro paragrafo."
        assert records[1]["text"] == "Segundo paragrafo."

    @logged
    def test_ignores_blank_paragraphs(self):
        content = b"Um.\n\n\n\nDois."
        records = extract(content, "x.txt")
        assert len(records) == 2

    @logged
    def test_single_paragraph(self):
        content = b"Apenas um paragrafo."
        records = extract(content, "notes.txt")
        assert records == [{"text": "Apenas um paragrafo."}]

    @logged
    def test_handles_utf8_bom(self):
        """Arquivo .txt com BOM UTF-8 deve ser decodificado corretamente."""
        bom_content = "﻿Como vai?".encode("utf-8")
        records = extract(bom_content, "bom.txt")
        assert records == [{"text": "Como vai?"}]


@pytest.mark.unit
@pytest.mark.part1
class TestExtractCsv:
    """Arquivos .csv — DictReader com header na primeira linha."""

    @logged
    def test_returns_one_dict_per_row(self):
        csv_content = b"nome,idade\nAlice,30\nBob,25\n"
        records = extract(csv_content, "people.csv")
        assert len(records) == 2

    @logged
    def test_uses_header_row_as_keys(self):
        csv_content = b"produto,preco\nWidget,10.50\n"
        records = extract(csv_content, "items.csv")
        assert records[0] == {"produto": "Widget", "preco": "10.50"}

    @logged
    def test_empty_csv_returns_empty_list(self):
        records = extract(b"col1,col2\n", "vazio.csv")
        assert records == []


@pytest.mark.unit
@pytest.mark.part1
class TestExtractJson:
    """Arquivos .json — lista é mantida, dict vira lista de 1 item."""

    @logged
    def test_list_of_objects_returned_as_is(self):
        data = [{"a": 1}, {"b": 2}]
        records = extract(json.dumps(data).encode(), "data.json")
        assert records == data

    @logged
    def test_single_object_wrapped_in_list(self):
        data = {"key": "value"}
        records = extract(json.dumps(data).encode(), "obj.json")
        assert records == [data]

    @logged
    def test_list_of_primitives_wraps_in_value_key(self):
        """Lista de strings/ints vira lista de {'value': ...}."""
        data = ["um", "dois", "tres"]
        records = extract(json.dumps(data).encode(), "list.json")
        assert records == [{"value": "um"}, {"value": "dois"}, {"value": "tres"}]


@pytest.mark.unit
@pytest.mark.part1
class TestExtractDocx:
    """Arquivos .docx — extração de parágrafos via python-docx."""

    @logged
    def test_extracts_paragraphs(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Primeiro paragrafo.")
        doc.add_paragraph("Segundo paragrafo.")
        buf = io.BytesIO()
        doc.save(buf)
        records = extract(buf.getvalue(), "relatorio.docx")
        texts = [r["text"] for r in records]
        assert "Primeiro paragrafo." in texts
        assert "Segundo paragrafo." in texts

    @logged
    def test_ignores_empty_paragraphs(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Conteudo.")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        buf = io.BytesIO()
        doc.save(buf)
        records = extract(buf.getvalue(), "x.docx")
        assert all(r["text"].strip() for r in records)


@pytest.mark.unit
@pytest.mark.part1
class TestExtractExcel:
    """Arquivos .xlsx — primeira linha é header, demais são linhas de dados."""

    @logged
    def test_extracts_rows_with_header(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["nome", "idade"])
        ws.append(["Alice", 30])
        ws.append(["Bob", 25])
        buf = io.BytesIO()
        wb.save(buf)
        records = extract(buf.getvalue(), "people.xlsx")
        assert len(records) == 2
        assert records[0]["nome"] == "Alice"
        assert records[0]["idade"] == 30

    @logged
    def test_ignores_completely_empty_rows(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["col"])
        ws.append(["valor"])
        ws.append([None])  # linha totalmente vazia
        buf = io.BytesIO()
        wb.save(buf)
        records = extract(buf.getvalue(), "x.xlsx")
        assert len(records) == 1


@pytest.mark.unit
@pytest.mark.regression
@pytest.mark.part1
class TestExtractUnsupported:
    """Regressão: extensão não suportada DEVE levantar ValueError com mensagem clara."""

    @logged
    def test_pptx_raises_value_error(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract(b"qualquer-coisa", "slides.pptx")

    @logged
    def test_unknown_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            extract(b"data", "arquivo.xyz")
